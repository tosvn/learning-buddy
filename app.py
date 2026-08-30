import os
import json
import uuid
import PyPDF2
import docx
from flask import Flask, render_template, request, jsonify, session
from groq import Groq
from werkzeug.utils import secure_filename

app = Flask(__name__)
app.secret_key = "chatbot_secret_key_2025"

# Groq config
client = Groq(api_key=os.environ.get("GROQ_API_KEY", "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"))

UPLOAD_FOLDER = "uploads"
HISTORY_FOLDER = "histories"
ALLOWED_EXTENSIONS = {"pdf", "docx"}
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(HISTORY_FOLDER, exist_ok=True)

SYSTEM_PROMPT = """You are Learning Buddy, an intelligent AI-based course assistant for undergraduate students at Yakubu Gowon University. Your role is to help students with academic questions, explain concepts clearly, review and summarise uploaded documents, assist with exam preparation, provide feedback on written work, and offer helpful academic guidance across all university courses and disciplines. 

You should ONLY respond to academic and educational requests. If a student asks about something completely unrelated to academics — such as entertainment, sports gossip, romantic topics, or personal matters — politely decline and remind them that you are designed for academic support only. However, always assist with any genuine study-related request including document review, exam preparation, essay feedback, concept explanation, and course-related questions."""

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(filepath):
    text = ""
    with open(filepath, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            text += page.extract_text() or ""
    return text

def extract_text_from_docx(filepath):
    doc = docx.Document(filepath)
    return "\n".join([para.text for para in doc.paragraphs])

def get_history_path(session_id):
    return os.path.join(HISTORY_FOLDER, f"{session_id}.json")

def load_history(session_id):
    path = get_history_path(session_id)
    if os.path.exists(path):
        with open(path, "r") as f:
            return json.load(f)
    return []

def save_history(session_id, messages):
    os.makedirs(HISTORY_FOLDER, exist_ok=True)
    path = get_history_path(session_id)
    with open(path, "w") as f:
        json.dump(messages, f)
        
def query_groq(messages):
    response = client.chat.completions.create(
       model="openai/gpt-oss-20b",
        messages=[{"role": "system", "content": SYSTEM_PROMPT}] + messages
    )
    return response.choices[0].message.content

@app.route("/")
def index():
    if "session_id" not in session:
        session["session_id"] = str(uuid.uuid4())
    save_history(session["session_id"], [])
    return render_template("index.html")

@app.route("/chat", methods=["POST"])
def chat():
    data = request.get_json()
    user_message = data.get("message", "")
    session_id = session.get("session_id", str(uuid.uuid4()))

    messages = load_history(session_id)
    messages.append({"role": "user", "content": user_message})

    reply = query_groq(messages)
    messages.append({"role": "assistant", "content": reply})
    save_history(session_id, messages)

    return jsonify({"reply": reply})

@app.route("/upload", methods=["POST"])
def upload():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]
    if file.filename == "" or not allowed_file(file.filename):
        return jsonify({"error": "Invalid file type. Please upload a PDF or Word document."}), 400

    filename = secure_filename(file.filename)
    
    # Process file directly from memory without saving to disk
    try:
        if filename.endswith(".pdf"):
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
        else:
            doc = docx.Document(file)
            text = "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        return jsonify({"error": f"Could not process document: {str(e)}"}), 400

    if not text.strip():
        return jsonify({"error": "Could not extract text from the document."}), 400

    session_id = session.get("session_id", str(uuid.uuid4()))
    messages = load_history(session_id)

    doc_message = f"The student has uploaded a document titled '{filename}'. Here is its content:\n\n{text[:3000]}"
    messages.append({"role": "user", "content": doc_message})
    messages.append({"role": "assistant", "content": f"I have received and reviewed '{filename}'. Please go ahead and ask me any questions about it."})
    save_history(session_id, messages)

    return jsonify({"message": f"Document '{filename}' uploaded successfully! You can now ask questions about it."})

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)), debug=False)